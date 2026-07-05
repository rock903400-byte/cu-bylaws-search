import os
import sqlite3
import fitz  # PyMuPDF
from docx import Document
from google.cloud import storage
from dotenv import load_dotenv

# 載入 .env 檔案
load_dotenv()

# 設定
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_DIR = os.path.join(BASE_DIR, "data")
DB_PATH = os.path.join(BASE_DIR, "legal_index.db")

# GCS 設定
GCS_BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME", "my-legal")
GCS_INTERNAL_BUCKET_NAME = os.environ.get("GCS_INTERNAL_BUCKET_NAME", "my-legal2")


def get_category(filename):
    categories = {
        "法": "法律",
        "辦法": "辦法",
        "要點": "要點",
        "規則": "規則",
        "章程": "章程",
        "簡則": "簡則",
        "須知": "作業須知",
        "手冊": "手冊",
        "範例": "範例",
    }
    for kw, cat in categories.items():
        if kw in filename:
            return cat
    return "其他"


def extract_text_pdf(path):
    try:
        doc = fitz.open(path)
        text = "".join([page.get_text() for page in doc])
        doc.close()
        return text
    except Exception as e:
        print(f"Error reading PDF {path}: {e}")
        return ""


def extract_text_docx(path):
    """讀取 DOCX 檔案，包含段落與所有表格內容"""
    try:
        doc = Document(path)
        full_text = []

        # 1. 讀取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # 2. 讀取所有表格中的儲存格內容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 去除儲存格內的換行，並加入列表
                    clean_cell = cell.text.strip().replace("\n", " ")
                    if clean_cell:
                        full_text.append(clean_cell)

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX {path}: {e}")
        return ""


def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DROP TABLE IF EXISTS docs")
    cursor.execute("""
        CREATE VIRTUAL TABLE docs USING fts5(
            title, category, content, path,
            filename UNINDEXED, source UNINDEXED,
            tokenize = 'trigram'
        )
    """)
    conn.commit()
    return conn


def index_source(conn, source_name):
    """索引特定來源的所有檔案 (包含子目錄與根目錄)"""
    cursor = conn.cursor()

    # 收集所有可能的目錄
    dirs_to_scan = []
    if source_name == "internal":
        dirs_to_scan.extend(
            [os.path.join(SOURCE_DIR, "internal"), os.path.join(SOURCE_DIR, "內部法規")]
        )
    else:
        dirs_to_scan.extend([os.path.join(SOURCE_DIR, "association"), SOURCE_DIR])

    all_files = []
    processed_filenames = set()  # 避免重複索引

    for d in dirs_to_scan:
        if os.path.exists(d):
            try:
                files = [
                    f
                    for f in os.listdir(d)
                    if f.lower().endswith((".pdf", ".docx"))
                    and os.path.isfile(os.path.join(d, f))
                ]
                for f in files:
                    if f not in processed_filenames:
                        all_files.append((d, f))
                        processed_filenames.add(f)
            except Exception as e:
                print(f"Error reading {d}: {e}")

    if not all_files:
        print(f"No files found for {source_name}")
        return

    print(f"Found {len(all_files)} total files to index for {source_name}.")

    indexed_count = 0
    for file_dir, filename in all_files:
        path = os.path.join(file_dir, filename)
        category = get_category(filename)
        title = os.path.splitext(filename)[0]

        print(
            f"[{source_name}] [{indexed_count+1}/{len(all_files)}] Indexing: {filename}"
        )

        try:
            if filename.lower().endswith(".pdf"):
                content = extract_text_pdf(path)
            else:
                content = extract_text_docx(path)

            # 儲存相對路徑以便前端讀取
            rel_path = f"{source_name}/{filename}"

            cursor.execute(
                "INSERT INTO docs (title, category, content, path, filename, source) VALUES (?, ?, ?, ?, ?, ?)",
                (title, category, content, rel_path, filename, source_name),
            )
            indexed_count += 1
        except Exception as e:
            print(f"Skipping {filename}: {e}")
    conn.commit()


def upload_db_to_gcs():
    if not GCS_BUCKET_NAME:
        return
    print(f"--- Uploading {DB_PATH} to GCS bucket: {GCS_BUCKET_NAME} ---")
    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(os.path.basename(DB_PATH))
        blob.upload_from_filename(DB_PATH)
        print("Database upload successful!")
    except Exception as e:
        print(f"Error uploading DB: {e}")


def run_indexer():
    print("--- Starting indexer ---")
    if not os.path.exists(SOURCE_DIR):
        os.makedirs(SOURCE_DIR)
    conn = init_db()
    index_source(conn, "association")
    index_source(conn, "internal")
    conn.close()
    upload_db_to_gcs()
    print("--- All Indexing complete! ---")


if __name__ == "__main__":
    run_indexer()
